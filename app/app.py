import os
import io
import tempfile
import traceback
import whisper
import requests
import bcrypt
from fpdf import FPDF
from docx import Document
from datetime import datetime
from flask import Flask, request, jsonify, session, redirect, render_template, make_response, send_file
from flask_cors import CORS
from flask_jwt_extended import JWTManager, create_access_token
from pymongo import MongoClient
from bson.objectid import ObjectId
from transformers import pipeline
from moviepy.editor import VideoFileClip
from pyannote.audio import Pipeline as DiarizationPipeline
from cryptography.fernet import Fernet
import re
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
from sentence_transformers import SentenceTransformer
import threading
from itsdangerous import URLSafeTimedSerializer, BadSignature, SignatureExpired
from flask import url_for, flash  
from flask import render_template, request, redirect, url_for, flash
from itsdangerous import URLSafeTimedSerializer, BadSignature, SignatureExpired
from flask_mail import Mail, Message
import time
import torch
from transformers import AutoTokenizer, AutoModelForCausalLM, TextIteratorStreamer
from peft import PeftModel

# Base model path (Phi-3 Mini by Microsoft)
BASE_MODEL_PATH = os.getenv("BASE_MODEL_PATH", "microsoft/phi-3-mini-4k-instruct") 

# LoRA adapter (fine-tuned weights for chatbot)
LORA_ADAPTER_PATH = "C:/Users/ADMIN/OneDrive/Desktop/Final MeetingSum working/buddhini research/app/pretend-lora-phi3-mini" #adapt to your path, fine-tuned lora adapter path

# Load model in 4-bit precision (saves VRAM, fits smaller GPUs ~4GB)
LOAD_4BIT_KW = dict(
    load_in_4bit=True,
    bnb_4bit_quant_type="nf4",
    bnb_4bit_compute_dtype=torch.float16,
    bnb_4bit_use_double_quant=True,
    device_map="auto"
)

print(f"Loading base model: {BASE_MODEL_PATH}")
_local_tokenizer = AutoTokenizer.from_pretrained(BASE_MODEL_PATH)
_local_model = AutoModelForCausalLM.from_pretrained(BASE_MODEL_PATH, **LOAD_4BIT_KW)

# Apply LoRA fine-tuned adapter if provided
if LORA_ADAPTER_PATH:
    print(f"Loading LoRA adapter: {LORA_ADAPTER_PATH}")
    _local_model = PeftModel.from_pretrained(_local_model, LORA_ADAPTER_PATH)
    _local_model.eval()

# Load encryption key
with open("fernet.key", "rb") as key_file:
    fernet = Fernet(key_file.read())

# MongoDB setup
client = MongoClient("mongodb+srv://minutesai_user:your_secure_password@minutesai.vs5njhc.mongodb.net/?retryWrites=true&w=majority")
db = client["minutesai"]

# Collections
users_collection = db["users"]  # Stores user accounts
transcripts_collection = db["transcripts"] # Stores meeting uploads
chat_hist = db["chat_history"]  # Stores chatbot Q&A history
chunks_collection = db["transcript_chunks"] # Stores transcript chunks (embeddings)

# Indexes for speed
chat_hist.create_index([("user_email", 1), ("doc_id", 1)])
chunks_collection.create_index([("doc_id", 1)])

# SentenceTransformer for semantic search inside transcripts
embedder = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")


app = Flask(__name__, template_folder='../templates', static_folder="../static")
CORS(app)

app.secret_key = "supersecretkey"

# Limit uploads to 100 MB
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100 MB

# Gmail SMTP config for sending password reset emails
app.config.update(
    MAIL_SERVER='smtp.gmail.com',
    MAIL_PORT=587,
    MAIL_USE_TLS=True,
    MAIL_USERNAME='buddhinialwis1@gmail.com',   
    MAIL_PASSWORD='gwug mqpx zuwy nqmp',      
    MAIL_DEFAULT_SENDER=('MinutesAI', 'yourgmail@gmail.com')
)

# Initialize Flask-Mail
mail = Mail(app)
ts = URLSafeTimedSerializer(app.secret_key)

from werkzeug.exceptions import RequestEntityTooLarge

@app.errorhandler(RequestEntityTooLarge)
def handle_file_too_large(e):
    return jsonify({'error': 'File too large. Max upload size is 100 MB.'}), 413

def make_reset_token(email: str) -> str:
    # short, signed token bound to your secret key
    return ts.dumps(email, salt="password-reset")

def get_email_from_token(token: str, max_age_seconds: int = 600) -> str:
    # 1 hour expiry by default
    return ts.loads(token, salt="password-reset", max_age=max_age_seconds)

# Ensure upload folder exists
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

#Whisper and BART models selected
model = whisper.load_model("tiny").to("cuda")
summarizer = pipeline("summarization", model="facebook/bart-large-cnn")


# Utility Functions
def encrypt(text):
    return fernet.encrypt(text.encode()).decode()
def decrypt(text):
    return fernet.decrypt(text.encode()).decode()

#define the function that extracts audio from video using moviepy
def extract_audio_from_video(video_path):
    try:
        video_clip = VideoFileClip(video_path)
        audio_clip = video_clip.audio
        temp_audio_file = os.path.join(UPLOAD_FOLDER, "temp_audio.wav")
        audio_clip.write_audiofile(temp_audio_file)
        return temp_audio_file
    except Exception as e:
        traceback.print_exc()
        return None
    
#define the function that extracts text from audio using Whisper
def extract_text_from_audio(audio_path):
    try:
        result = model.transcribe(audio_path)
        return result["text"].strip(), result["segments"]
    except:
        return None, []
    

def diarize_audio(audio_path):
    """Separate speakers using pyannote diarization model."""
    diarization_pipeline = DiarizationPipeline.from_pretrained(
        "pyannote/speaker-diarization@2.1",
        use_auth_token="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"
    )
    return diarization_pipeline(audio_path)

def combine_diarization_with_transcription(diarization, segments):
    """Merge Whisper transcription segments with speaker diarization."""
    output = ""
    for turn, _, speaker in diarization.itertracks(yield_label=True):
        text = ""
        for seg in segments:
            if not (seg['end'] < turn.start or seg['start'] > turn.end):
                text += seg['text'].strip() + " "
        start_time = f"{int(turn.start // 60):02d}:{int(turn.start % 60):02d}"
        end_time = f"{int(turn.end // 60):02d}:{int(turn.end % 60):02d}"
        output += f"Speaker {speaker} ({start_time}-{end_time}): {text.strip()}\n\n"
    return output

# simple BART summarization with length limits and error handling
def summarize_text(text):
    try:
        max_words = 1024
        if len(text.split()) > max_words:
            text = " ".join(text.split()[:max_words])
        summary = summarizer(text, max_length=130, min_length=30, do_sample=False)
        return summary[0]['summary_text']
    except:
        return "Failed to summarize"

#calculate word error rate
def calculate_wer(reference: str, hypothesis: str) -> float:
    ref_words = reference.strip().split()
    hyp_words = hypothesis.strip().split()
    n = len(ref_words)

    dp = np.zeros((len(ref_words)+1, len(hyp_words)+1), dtype=np.int32)
    for i in range(len(ref_words)+1):
        dp[i][0] = i
    for j in range(len(hyp_words)+1):
        dp[0][j] = j

    for i in range(1, len(ref_words)+1):
        for j in range(1, len(hyp_words)+1):
            if ref_words[i-1] == hyp_words[j-1]:
                dp[i][j] = dp[i-1][j-1]
            else:
                substitution = dp[i-1][j-1] + 1
                insertion = dp[i][j-1] + 1
                deletion = dp[i-1][j] + 1
                dp[i][j] = min(substitution, insertion, deletion)

    wer_score = dp[len(ref_words)][len(hyp_words)] / float(n) if n > 0 else 0
    wer_score = min(wer_score, 1.0) 
    return round(wer_score * 100, 2)

# -----------------------------
# CHATBOT MEMORY + RAG FUNCTIONS
# -----------------------------
# (get_history, save_turn, build_messages, _messages_to_prompt, parse_diarized, 
# make_chunks, embed_and_store, retrieve, local_llm_chat)
# → These functions:
#    - Store/retrieve chat history
#    - Split transcripts into chunks + embed them
#    - Retrieve top relevant chunks
#    - Build context + prompt for LLM
#    - Generate answer with local Phi-3 Mini model


from datetime import datetime

def get_history(user_email, doc_id, limit=8):
    doc = chat_hist.find_one({"user_email": user_email, "doc_id": str(doc_id)})
    turns = doc["turns"] if doc else []
    return turns[-limit:]

def save_turn(user_email, doc_id, role, text):
    chat_hist.update_one(
        {"user_email": user_email, "doc_id": str(doc_id)},
        {"$push": {"turns": {"role": role, "text": text, "t": datetime.utcnow()}}},
        upsert=True
    )
# Build ChatGPT-style messages for the local LLM
def build_messages(system_rules, history, retrieved_snips, question):
    # Conversational style + safety
    system_prompt = (
        "You are a helpful meeting assistant.\n"
        "Speak concisely and conversationally, like ChatGPT.\n"
        "Answer ONLY using the provided meeting context; if it's not there, say you don't know.\n"
        "When you cite evidence, include timestamp ranges like [00:10–00:24].\n"
        "Prefer direct answers first, then bullets if needed. Avoid hallucinations."
    )
    if system_rules:
        system_prompt = system_rules + "\n\n" + system_prompt

    # Compact context text
    ctx = []
    for s in retrieved_snips:
        ctx.append(f"[{s['start']}–{s['end']}]\n{s['text']}")
    ctx_text = "\n\n".join(ctx)

    messages = [{"role": "system", "content": system_prompt}]
    # Short multi-turn memory (last 6–8 turns)
    for turn in history:
        messages.append({"role": turn["role"], "content": turn["text"]})

    # Add context as an assistant “context drop”
    messages.append({
        "role": "assistant",
        "content": "Here is the relevant meeting context (use ONLY this to answer):\n\n" + ctx_text
    })
    # User question last
    messages.append({"role": "user", "content": question})
    return messages
time_pat = re.compile(r"Speaker\s+(?P<spk>\S+)\s+\((?P<start>\d\d:\d\d)-(?P<end>\d\d:\d\d)\):\s*(?P<text>.+)")

def _messages_to_prompt(messages):
    """
    Convert ChatGPT-style messages into a single prompt for an instruct model.
    Keep it consistent with how you fine-tuned (Question/Input/Answer).
    """
    sys_text = ""
    conv_lines = []
    for m in messages:
        role = m["role"]
        content = m["content"]
        if role == "system":
            sys_text = content.strip()
        elif role == "user":
            conv_lines.append(f"User: {content.strip()}")
        elif role == "assistant":
            # We inject retrieved context as assistant preamble; mark it clearly
            conv_lines.append(f"Assistant (context): {content.strip()}")
        else:
            conv_lines.append(f"{role.capitalize()}: {content.strip()}")

    history_block = "\n".join(conv_lines)
    sys_block = f"System: {sys_text}\n" if sys_text else ""
    # Simple, instruction-aligned template
    prompt = (
        f"{sys_block}"
        f"{history_block}\n\n"
        f"Assistant: "
    )
    return prompt

# Embedding + Retrieval for Chatbot
def parse_diarized(diarized_text: str):
    out = []
    for line in diarized_text.strip().splitlines():
        m = time_pat.match(line.strip())
        if m:
            out.append({
                "speaker": m.group("spk"),
                "start": m.group("start"),
                "end":   m.group("end"),
                "text":  m.group("text").strip()
            })
    return out
# Chunking: max 800 chars with 200 char overlap
def make_chunks(utterances, max_chars=800, overlap_chars=200):
    chunks = []; buf=[]; size=0; start=None
    def flush():
        nonlocal buf, size, start
        if buf:
            text = "\n".join(buf)
            # get end from last line’s [mm:ss–mm:ss]
            m = re.search(r"\[(\d\d:\d\d)–(\d\d:\d\d)\]", buf[-1])
            end = m.group(2) if m else start
            chunks.append({"text": text, "start": start, "end": end})
            buf=[]; size=0; start=None

    for u in utterances:
        seg = f"[{u['start']}–{u['end']}] {u['speaker']}: {u['text']}"
        if not buf: start = u["start"]
        if size + len(seg) > max_chars and buf:
            flush()
            # overlap: keep tail of previous buf
            # (simple version—good enough)
            pass
        buf.append(seg); size += len(seg)
    flush()
    return chunks
# Embed + store chunks for a given transcript
def embed_and_store(doc_id, diarized_text):
    utter = parse_diarized(diarized_text)
    if not utter:
        return
    chunks = make_chunks(utter, max_chars=900, overlap_chars=200)
    texts = [c["text"] for c in chunks]
    vecs = embedder.encode(texts, convert_to_numpy=True, normalize_embeddings=True)
    chunks_collection.delete_many({"doc_id": str(doc_id)})
    recs = [{
        "doc_id": str(doc_id),
        "text": c["text"],
        "start": c["start"],
        "end": c["end"],
        "embedding": vecs[i].astype(np.float32).tolist()
    } for i, c in enumerate(chunks)]
    if recs:
        chunks_collection.insert_many(recs)

# Retrieve top-k relevant chunks for a question
def retrieve(doc_id, question, top_k=4):
    qv = embedder.encode([question], convert_to_numpy=True, normalize_embeddings=True)[0]
    docs = list(chunks_collection.find({"doc_id": str(doc_id)}, {"_id":0}))
    if not docs:
        return []
    M = np.array([d["embedding"] for d in docs], dtype=np.float32)
    sims = (M @ qv)  # cosine: rows are normalized
    idx = np.argsort(-sims)[:top_k]
    return [docs[i] | {"score": float(sims[i])} for i in idx]

# Local LLM inference function
def local_llm_chat(messages, max_new_tokens=256, temperature=0.2, top_p=0.9):
    prompt = _messages_to_prompt(messages)

    # Tokenize
    inputs = _local_tokenizer(
        prompt,
        return_tensors="pt",
        truncation=True,
        max_length=2048
    )
    inputs = {k: v.to(_local_model.device) for k, v in inputs.items()}

    with torch.no_grad():
        output_ids = _local_model.generate(
            **inputs,
            max_new_tokens=max_new_tokens,
            do_sample=(temperature > 0),
            temperature=temperature,
            top_p=top_p,
            repetition_penalty=1.1,
            pad_token_id=_local_tokenizer.eos_token_id,
            eos_token_id=_local_tokenizer.eos_token_id
        )
    full_text = _local_tokenizer.decode(output_ids[0], skip_special_tokens=True)

    # Return only the part after the last "Assistant: "
    sep = "Assistant:"
    if sep in full_text:
        return full_text.split(sep)[-1].strip()
    return full_text.strip()

#define the chatbot endpoint
@app.post("/ask_chatbot")
def ask_chatbot():
    if "user_email" not in session:
        return jsonify({"response": "Please log in."}), 401

    data = request.get_json(force=True)
    doc_id = data.get("doc_id")
    question = (data.get("question") or "").strip()
    if not doc_id or not question:
        return jsonify({"response": "Missing doc_id or question."}), 400

    doc = transcripts_collection.find_one({"_id": ObjectId(doc_id), "user_email": session["user_email"]})
    if not doc:
        return jsonify({"response": "Transcript not found."}), 404

    # Make sure chunks exist (index on first chat or after upload)
    if chunks_collection.count_documents({"doc_id": str(doc_id)}) == 0:
        diarized_text = decrypt(doc["diarized_text"])
        embed_and_store(doc_id, diarized_text)

    # Multi-turn history
    history = get_history(session["user_email"], doc_id, limit=8)

    # Retrieval: keep it small on 4 GB
    top = retrieve(doc_id, question, top_k=4)
    if not top:
        return jsonify({"response": "I couldn’t find anything relevant in this meeting."})

    # Build messages (ChatGPT-style)
    messages = build_messages(system_rules="", history=history, retrieved_snips=top, question=question)

    # Call small model for fluent answer
    try:
        answer_text = local_llm_chat(messages, temperature=0.2, max_new_tokens=220)
    except Exception as e:
        # Graceful fallback: stitch top snippets
        stitched = "\n\n".join([f"[{t['start']}–{t['end']}]\n{t['text']}" for t in top])
        answer_text = f"Here’s what I can confirm from the meeting:\n\n{stitched}\n\n(Note: generator unavailable: {e})"

    # Add inline timestamp refs at the end (simple but useful)
    refs = ", ".join([f"{t['start']}–{t['end']}" for t in top])
    final_answer = f"{answer_text}\n\nReferences: {refs}"

    # Save turns
    save_turn(session["user_email"], doc_id, "user", question)
    save_turn(session["user_email"], doc_id, "assistant", final_answer)

    return jsonify({"response": final_answer})

#define the home route
@app.route("/")
def home():
    return render_template("index.html")

#define the forgot password route
@app.route("/forgot-password", methods=["GET", "POST"])
def forgot_password():
    if request.method == "POST":
        submitted = (request.form.get("email") or "").strip().lower()

        # Find the relevant user
        user = users_collection.find_one({"email": submitted})

        if user:
            # Optional: basic throttle to stop abuse (1 email / 3 min per account)
            last = user.get("last_reset_sent_at", 0)
            now = int(time.time())
            if now - int(last) > 180:  # 180s = 3 minutes
                token = make_reset_token(submitted)
                reset_link = url_for("reset_password", token=token, _external=True)

                # Personalize
                display_name = user.get("name") or "there"

                try:
                    msg = Message(
                        subject="Reset your MinutesAI password",
                        recipients=[submitted]  # ← send only to the account email
                    )
                    # Plain text (helps deliverability)
                    msg.body = (
                        f"Hi {display_name},\n\n"
                        f"Use the link below to reset your MinutesAI password. "
                        f"This link expires in 10 minutes.\n\n{reset_link}\n\n"
                        "If you didn’t request this, you can ignore this email."
                    )
                    # HTML version
                    msg.html = f"""
                      <p>Hi {display_name},</p>
                      <p>Click the button below to reset your MinutesAI password (valid for 10 minutes).</p>
                      <p><a href="{reset_link}" style="background:#4f46e5;color:#fff;padding:10px 16px;border-radius:8px;text-decoration:none;">Reset Password</a></p>
                      <p>If the button doesn’t work, copy this link into your browser:<br>
                      <code>{reset_link}</code></p>
                      <p>If you didn’t request this, you can safely ignore this email.</p>
                    """
                    mail.send(msg)

                    # Save throttle timestamp
                    users_collection.update_one(
                        {"_id": user["_id"]},
                        {"$set": {"last_reset_sent_at": now}}
                    )
                except Exception:
                    # Don’t leak details; optionally log server-side
                    pass

        # Same response either way (prevents email enumeration)
        flash("If that email exists, a reset link has been sent.", "info")
        return redirect(url_for("login"))

    return render_template("forgot_password.html")

#define the reset password route
@app.route("/reset-password/<token>", methods=["GET", "POST"])
def reset_password(token):
    try:
        email = get_email_from_token(token, max_age_seconds=3600)
    except (BadSignature, SignatureExpired):
        flash("The reset link is invalid or has expired.", "danger")
        return redirect(url_for("forgot_password"))

    if request.method == "POST":
        pwd = (request.form.get("password") or "").strip()
        confirm = (request.form.get("confirm_password") or "").strip()
        if len(pwd) < 8:
            flash("Password must be at least 8 characters.", "warning"); return redirect(request.url)
        if pwd != confirm:
            flash("Passwords do not match.", "warning"); return redirect(request.url)

        hashed = bcrypt.hashpw(pwd.encode("utf-8"), bcrypt.gensalt())
        users_collection.update_one({"email": email}, {"$set": {"password": hashed}})
        session.clear()
        flash("Your password has been updated. Please log in.", "success")
        return redirect(url_for("login"))

    return render_template("reset_password.html")

#define the signup route
@app.route("/signup", methods=["GET", "POST"])
def signup():
    if request.method == "POST":
        name = request.form["name"]
        email = request.form["email"].lower()
        password = request.form["password"].encode('utf-8')
        if users_collection.find_one({"email": email}):
            return "User already exists. Try logging in."
        hashed = bcrypt.hashpw(password, bcrypt.gensalt())
        users_collection.insert_one({"name": name, "email": email, "password": hashed})
        return redirect("/login")
    return render_template("signup.html")

#define the login route
@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = (request.form.get("email") or "").strip().lower()
        password = (request.form.get("password") or "")
        user = users_collection.find_one({"email": email})

        if not user:
            # same message for security
            return render_template("login.html", error="Invalid email or password.")

        # ensure both are bytes
        pw_bytes = password.encode("utf-8")
        stored = user.get("password")
        if isinstance(stored, str):
            stored = stored.encode("utf-8")
        else:
            try:
                # convert BSON Binary to raw bytes if needed
                stored = bytes(stored)
            except Exception:
                pass

        if not bcrypt.checkpw(pw_bytes, stored):
            return render_template("login.html", error="Invalid email or password.")

        # success: set session and go
        session["user_email"] = email
        session["user_name"]  = user.get("name", "User")
        return redirect("/dashboard")

    return render_template("login.html")

@app.route("/logout")
def logout():
    session.clear()
    return redirect("/login")

@app.route("/dashboard")
def dashboard():
    if "user_name" not in session:
        return redirect("/login")
    return render_template("dashboard.html", username=session["user_name"])

#define the upload media route
@app.route("/upload-media", methods=["POST"])
def upload_media():
    """
    Handles file upload, transcription, diarization, summarization, and storage.
    Also triggers background embedding + indexing for chatbot retrieval.
    """
    file = request.files.get("file")
    reference_text = request.form.get("reference_text", "").strip()
    title = request.form.get("title", "").strip()
    date = request.form.get("date", "").strip()

    if not file or file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    # Save uploaded file
    file_path = os.path.join(UPLOAD_FOLDER, file.filename)
    file.save(file_path)

    try:
        # Extract audio from video if needed
        audio_path = extract_audio_from_video(file_path) if file.filename.endswith((".mp4", ".webm", ".avi")) else file_path

        # Transcribe audio
        transcribed_text, segments = extract_text_from_audio(audio_path)

        # Speaker diarization
        diarization = diarize_audio(audio_path)
        diarized_text = combine_diarization_with_transcription(diarization, segments)

        # Generate summary
        summary = summarize_text(transcribed_text)

        # Cleanup temporary audio if it was extracted from video
        if audio_path != file_path:
            os.remove(audio_path)

        # Store transcript in MongoDB
        result = transcripts_collection.insert_one({
            "user_email": session.get("user_email", "guest"),
            "file_name": file.filename,
            "meeting_title": title,
            "meeting_date": date,
            "upload_date": datetime.utcnow(),
            "transcript": encrypt(transcribed_text),
            "summary": encrypt(summary),
            "diarized_text": encrypt(diarized_text)
        })

        # Trigger background embedding for chatbot retrieval
        try:
            threading.Thread(
                target=embed_and_store,
                args=(result.inserted_id, diarized_text),
                daemon=True
            ).start()
        except Exception as e:
            print("Failed to start background indexing:", e)

        # Prepare response
        res = {
            'message': 'File processed',
            'summary': summary,
            'diarized_text': diarized_text,
            'doc_id': str(result.inserted_id)
        }

        # If reference text provided, calculate accuracy
        if reference_text:
            res['WER'] = f"{calculate_wer(transcribed_text, reference_text)}%"

        return jsonify(res)

    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

#define the route to view user's transcripts
@app.route("/my_transcripts")
def my_transcripts():
    if "user_email" not in session:
        return redirect("/login")
    docs = transcripts_collection.find({"user_email": session["user_email"]}).sort("upload_date", -1)
    results = []
    for doc in docs:
        results.append({
            "id": str(doc["_id"]),
            "title": doc.get("meeting_title", "Untitled"),
            "date": doc.get("meeting_date", "Unknown"),
            "file": doc.get("file_name", ""),
            "summary": decrypt(doc.get("summary", "")),
            "transcript": decrypt(doc.get("diarized_text", ""))
        })
    return render_template("my_transcripts.html", username=session["user_name"], uploads=results)

#define the route to view a specific transcript
@app.route("/view_transcript/<doc_id>")
def view_transcript(doc_id):
    if "user_email" not in session:
        return redirect("/login")
    doc = transcripts_collection.find_one({"_id": ObjectId(doc_id), "user_email": session["user_email"]})
    if not doc:
        return "Transcript not found", 404
    return render_template("view_transcript.html",
        meeting_title=doc.get("meeting_title", "Untitled"),
        summary=decrypt(doc["summary"]),
        transcript=decrypt(doc["diarized_text"]),
        file_id=doc["_id"]
    )
#define the route to download transcript as PDF or DOCX
@app.route("/download_transcript")
def download_transcript():
    if "user_email" not in session:
        return redirect("/login")

    doc_id = request.args.get("doc_id")
    format = request.args.get("format", "pdf")
    doc = transcripts_collection.find_one({"_id": ObjectId(doc_id), "user_email": session["user_email"]})
    if not doc:
        return "Transcript not found", 404

    summary = decrypt(doc.get("summary", ""))
    transcript = decrypt(doc.get("diarized_text", ""))
    title = doc.get("meeting_title", "meeting")
    content = f"Meeting Summary:\n\n{summary}\n\n\nSpeaker Transcript:\n\n{transcript}"

    if format == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)
        for line in content.split("\n"):
            pdf.multi_cell(0, 10, line)
        output = io.BytesIO(pdf.output(dest='S').encode('latin1'))
        return send_file(output, as_attachment=True, download_name=f"{title.replace(' ', '_')}.pdf", mimetype='application/pdf')

    elif format == "docx":
        docx = Document()
        docx.add_heading("Meeting Summary", level=1)
        docx.add_paragraph(summary)
        docx.add_page_break()
        docx.add_heading("Speaker Transcript", level=1)
        docx.add_paragraph(transcript)
        output = io.BytesIO()
        docx.save(output)
        output.seek(0)
        return send_file(output, as_attachment=True, download_name=f"{title.replace(' ', '_')}.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    return "Unsupported format", 400
from flask import url_for

@app.route("/delete_transcript", methods=["POST"])
def delete_transcript():
    if "user_email" not in session:
        return redirect(url_for("login"))

    doc_id = request.form.get("doc_id")
    result = transcripts_collection.delete_one(
        {"_id": ObjectId(doc_id), "user_email": session["user_email"]}
    )

    if result.deleted_count == 1:
        return redirect(url_for("my_transcripts"))   # clean redirect
    else:
        return "Not found", 404


#define the route to view user's chat sessions
@app.route("/my_chats")
def my_chats():
    if "user_email" not in session:
        return redirect("/login")
    docs = transcripts_collection.find({"user_email": session["user_email"]}).sort("upload_date", -1)
    uploads = [{"id": str(d["_id"]), "title": d.get("meeting_title","Untitled"), "date": d.get("meeting_date","Unknown")} for d in docs]
    return render_template("my_chats.html", uploads=uploads)

#define the chat interface route
@app.route("/chat/<doc_id>")
def chat(doc_id):
    if "user_email" not in session:
        return redirect("/login")
    doc = transcripts_collection.find_one({"_id": ObjectId(doc_id), "user_email": session["user_email"]})
    if not doc:
        return "Transcript not found", 404
    return render_template(
        "chat.html",
        username=session.get("user_name","User"),
        meeting_title=doc.get("meeting_title","Untitled"),
        doc_id=str(doc["_id"])
    )

@app.route("/settings")
def settings():
    if "user_email" not in session:
        return redirect("/login")

    user = users_collection.find_one({"email": session["user_email"]})
    return render_template("settings.html", user=user)

@app.route("/settings/update_profile", methods=["POST"])
def update_profile():
    if "user_email" not in session:
        return redirect("/login")

    new_name = (request.form.get("name") or "").strip()
    new_email = (request.form.get("email") or "").strip().lower()

    # find the current user
    current_user = users_collection.find_one({"email": session["user_email"]})
    if not current_user:
        return redirect("/login")

    # prevent email conflicts
    exists = users_collection.find_one({"email": new_email, "_id": {"$ne": current_user["_id"]}})
    if exists:
        return render_template("settings.html", user=current_user, message="Email already in use.")

    # update user record
    users_collection.update_one(
        {"_id": current_user["_id"]},
        {"$set": {"name": new_name, "email": new_email}}
    )

    # update session (so sidebar/header uses new name)
    session["user_email"] = new_email
    session["user_name"] = new_name

    updated_user = users_collection.find_one({"_id": current_user["_id"]})
    return render_template("settings.html", user=updated_user, message="Profile updated successfully.")


@app.route("/settings/update_password", methods=["POST"])
def update_password():
    if "user_email" not in session:
        return redirect("/login")

    current = request.form.get("current_password")
    new = request.form.get("new_password")
    confirm = request.form.get("confirm_password")

    user = users_collection.find_one({"email": session["user_email"]})
    stored = user.get("password")

    # convert stored to bytes if needed
    if isinstance(stored, str):
        stored = stored.encode("utf-8")

    if not bcrypt.checkpw(current.encode("utf-8"), stored):
        return render_template("settings.html", user=user, message="Current password incorrect.")

    if new != confirm:
        return render_template("settings.html", user=user, message="New passwords do not match.")

    hashed = bcrypt.hashpw(new.encode("utf-8"), bcrypt.gensalt())
    users_collection.update_one({"email": user["email"]}, {"$set": {"password": hashed}})

    return render_template("settings.html", user=user, message="Password updated successfully.")

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
